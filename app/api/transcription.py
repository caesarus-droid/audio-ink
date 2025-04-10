from flask import Blueprint, request, jsonify, current_app, render_template
from werkzeug.utils import secure_filename
import os
from typing import Dict, Any
from app.models.transcription import Transcription
import whisper
import torch
import time
import docx
from docx.shared import Pt
from io import BytesIO
from pathlib import Path
import tempfile
from .. import db
from datetime import datetime

api = Blueprint('transcription_api', __name__)

ALLOWED_EXTENSIONS = {'mp3', 'wav', 'ogg', 'mp4', 'm4a', 'mpeg', 'webm'}

# Initialize Whisper model
MODEL = None
DEVICE = None

def get_device_info() -> tuple[str, bool]:
    """Get device information and check CUDA availability."""
    if torch.cuda.is_available():
        device = "cuda"
        gpu_name = torch.cuda.get_device_name(0)
        current_app.logger.info(f"Using GPU: {gpu_name}")
        return device, True
    else:
        device = "cpu"
        current_app.logger.warning("CUDA not available, using CPU")
        return device, False

def load_model() -> tuple[whisper.Whisper, str]:
    """Load the Whisper model if not already loaded."""
    global MODEL, DEVICE
    if MODEL is None:
        current_app.logger.info("Loading Whisper model...")
        device, gpu_available = get_device_info()
        DEVICE = device
        MODEL = whisper.load_model("large-v3", device=device)
        current_app.logger.info(f"Whisper model loaded successfully on {device}")
    return MODEL, DEVICE

def allowed_file(filename: str) -> bool:
    """Check if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def transcribe_audio(audio_path: str) -> Dict[str, Any]:
    """Transcribe audio file to text using Whisper."""
    try:
        model, device = load_model()
        current_app.logger.info(f"Starting transcription of {audio_path}")
        
        start_time = time.time()
        result = model.transcribe(
            audio_path,
            language="en",
            task="transcribe",
            verbose=False
        )
        end_time = time.time()
        
        processing_time = end_time - start_time
        current_app.logger.info(f"Transcription completed in {processing_time:.2f} seconds")
        
        return {
            "text": result["text"],
            "processing_time": processing_time,
            "device": device,
            "segments": result.get("segments", []),
            "language": result.get("language", "en")
        }
        
    except Exception as e:
        current_app.logger.error(f"Error transcribing audio file {audio_path}: {str(e)}")
        return {
            "error": f"Error transcribing audio file: {str(e)}",
            "text": "",
            "processing_time": 0,
            "device": device,
            "segments": [],
            "language": "en"
        }

def get_whisper_model(use_gpu=True):
    """Get Whisper model with appropriate device configuration"""
    try:
        if use_gpu and torch.cuda.is_available():
            device = "cuda"
            current_app.logger.info("Using GPU for transcription")
        else:
            device = "cpu"
            current_app.logger.info("Using CPU for transcription")
        
        model = whisper.load_model("large-v3", device=device)
        return model
    except Exception as e:
        current_app.logger.error(f"Error loading Whisper model: {e}")
        return None

@api.route('/transcribe', methods=['POST'])
def transcribe():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Get user preference for GPU usage
    use_gpu = request.form.get('use_gpu', 'true').lower() == 'true'
    
    try:
        # Save the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name
        
        # Get the model with appropriate device configuration
        model = get_whisper_model(use_gpu)
        if model is None:
            return jsonify({'error': 'Failed to load transcription model'}), 500
        
        # Perform transcription
        result = model.transcribe(temp_path)
        
        # Create transcription record
        transcription = Transcription(
            filename=file.filename,
            text=result['text'],
            created_at=datetime.utcnow()
        )
        db.session.add(transcription)
        db.session.commit()
        
        # Clean up temporary file
        os.unlink(temp_path)
        
        return jsonify({
            'id': transcription.id,
            'text': result['text'],
            'filename': file.filename,
            'created_at': transcription.created_at.isoformat()
        })
        
    except Exception as e:
        current_app.logger.error(f"Transcription error: {e}")
        return jsonify({'error': str(e)}), 500

@api.route('/transcriptions', methods=['GET'])
def get_transcriptions():
    transcriptions = Transcription.query.order_by(Transcription.created_at.desc()).all()
    return jsonify([{
        'id': t.id,
        'filename': t.filename,
        'text': t.text,
        'created_at': t.created_at.isoformat()
    } for t in transcriptions])

@api.route('/transcriptions/<int:id>', methods=['GET'])
def get_transcription(id):
    transcription = Transcription.query.get_or_404(id)
    return jsonify({
        'id': transcription.id,
        'filename': transcription.filename,
        'text': transcription.text,
        'created_at': transcription.created_at.isoformat()
    })

@api.route('/transcriptions', methods=['POST'])
def create_transcription():
    """Create a new transcription task."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if not file or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type'}), 400
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Create transcription record
        transcription = Transcription(
            storage_path=current_app.config['STORAGE_PATH'],
            file_name=filename,
            file_path=filepath
        )
        transcription.save_to_json()
        
        return jsonify({
            'id': transcription.id,
            'status': 'pending',
            'message': 'File uploaded successfully. Starting transcription...'
        })
    
    except Exception as e:
        return jsonify({'error': f'Error uploading file: {str(e)}'}), 500

@api.route('/transcriptions/<transcription_id>/process', methods=['POST'])
def process_transcription(transcription_id: str):
    """Process a pending transcription."""
    transcription = Transcription.find_by_id(current_app.config['STORAGE_PATH'], transcription_id)
    if not transcription:
        return jsonify({'error': 'Transcription not found'}), 404
    
    if transcription.status != 'pending':
        return jsonify({'error': 'Transcription is not pending'}), 400
    
    try:
        result = transcribe_audio(transcription.file_path)
        
        if 'error' in result:
            transcription.mark_error(result['error'])
        else:
            transcription.update_result(
                text=result['text'],
                device=result['device'],
                processing_time=result['processing_time'],
                segments=result['segments'],
                language=result['language']
            )
        
        # Clean up temporary file
        try:
            os.remove(transcription.file_path)
        except Exception as e:
            current_app.logger.warning(f"Error cleaning up temporary file: {str(e)}")
        
        return jsonify(transcription.to_dict())
    
    except Exception as e:
        transcription.mark_error(str(e))
        return jsonify({'error': f'Error processing transcription: {str(e)}'}), 500

@api.route('/transcriptions', methods=['GET'])
def list_transcriptions():
    """List all transcriptions with optional status filter."""
    status = request.args.get('status')
    
    if status == 'pending':
        transcriptions = Transcription.get_pending(current_app.config['STORAGE_PATH'])
    elif status == 'completed':
        transcriptions = Transcription.get_completed(current_app.config['STORAGE_PATH'])
    elif status == 'failed':
        transcriptions = Transcription.get_failed(current_app.config['STORAGE_PATH'])
    else:
        transcriptions = Transcription.load_from_json(current_app.config['STORAGE_PATH'])
    
    return jsonify([t.to_dict() for t in transcriptions])

@api.route('/transcriptions/<transcription_id>/view')
def view_transcription(transcription_id: str):
    """View and edit a transcription."""
    transcription = Transcription.find_by_id(current_app.config['STORAGE_PATH'], transcription_id)
    if not transcription:
        return jsonify({'error': 'Transcription not found'}), 404
    
    return render_template('transcript.html', transcription_id=transcription_id)

@api.route('/transcriptions/<transcription_id>/download', methods=['POST'])
def download_transcription(transcription_id: str):
    """Download the edited transcription as a Word document."""
    transcription = Transcription.find_by_id(current_app.config['STORAGE_PATH'], transcription_id)
    if not transcription:
        return jsonify({'error': 'Transcription not found'}), 404
    
    try:
        content = request.json.get('content', '')
        
        # Create a new Word document
        doc = docx.Document()
        
        # Add the content
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(12)
        
        # Save the document to a BytesIO object
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        return doc_stream.getvalue(), 200, {
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'Content-Disposition': f'attachment; filename=transcript_{transcription_id}.docx'
        }
    
    except Exception as e:
        current_app.logger.error(f"Error creating Word document: {str(e)}")
        return jsonify({'error': f'Error creating Word document: {str(e)}'}), 500

@api.route('/gpu-status', methods=['GET'])
def gpu_status():
    """Check if GPU is available for transcription"""
    try:
        gpu_available = torch.cuda.is_available()
        return jsonify({
            'gpu_available': gpu_available,
            'device_name': torch.cuda.get_device_name(0) if gpu_available else None
        })
    except Exception as e:
        current_app.logger.error(f"Error checking GPU status: {e}")
        return jsonify({
            'gpu_available': False,
            'error': str(e)
        }) 